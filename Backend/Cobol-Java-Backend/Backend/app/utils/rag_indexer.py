import json
import logging
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from ..config import logger, AZURE_CONFIG, output_dir
import PyPDF2
from docx import Document as DocxDocument

RAG_DIR = Path(output_dir) / "rag"
STANDARDS_RAG_DIR = Path(output_dir) / "standards-rag"

def get_embedding_client():
    """Initialize Azure OpenAI embedding client."""
    try:
        embed_model = AzureOpenAIEmbedding(
            model=AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
            deployment_name=AZURE_CONFIG["AZURE_OPENAI_EMBED_DEPLOYMENT"],
            api_key=AZURE_CONFIG["AZURE_OPENAI_EMBED_API_KEY"],
            azure_endpoint=AZURE_CONFIG["AZURE_OPENAI_EMBED_API_ENDPOINT"],
            api_version=AZURE_CONFIG["AZURE_OPENAI_EMBED_VERSION"],
        )
        logger.info("Azure OpenAI embedding client initialized successfully")
        return embed_model
    except Exception as e:
        logger.error(f"Failed to initialize embedding client: {str(e)}")
        raise

embedding_client = get_embedding_client()

class AzureOpenAIEmbeddingWrapper:
    """Wrapper for Azure OpenAI embeddings compatible with LangChain."""
    def __init__(self, azure_embedding_client):
        self.client = azure_embedding_client
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        try:
            embeddings = []
            for text in texts:
                embedding = self.client.get_text_embedding(text)
                embeddings.append(embedding)
            return embeddings
        except Exception as e:
            logger.error(f"Error embedding documents: {str(e)}")
            raise
    
    def embed_query(self, text: str) -> List[float]:
        try:
            return self.client.get_text_embedding(text)
        except Exception as e:
            logger.error(f"Error embedding query: {str(e)}")
            raise

embedding_wrapper = AzureOpenAIEmbeddingWrapper(embedding_client)

def test_embedding_service():
    """Test the embedding service."""
    try:
        test_text = "This is a test document for embedding."
        embedding = embedding_wrapper.embed_query(test_text)
        logger.info(f"Embedding test successful. Dimension: {len(embedding)}")
        return True
    except Exception as e:
        logger.error(f"Embedding test failed: {str(e)}")
        return False

def extract_text_from_file(file_path: Path) -> str:
    """Extract text from PDF, DOC, DOCX, or TXT files."""
    logger.info(f"Extracting text from file: {file_path}")
    if file_path.suffix.lower() not in [".pdf", ".doc", ".docx", ".txt"]:
        logger.warning(f"Unsupported file type for {file_path}. Expected .pdf, .doc, .docx, or .txt.")
        return ""
    
    try:
        if file_path.suffix.lower() == ".pdf":
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted
                return text
        elif file_path.suffix.lower() in [".doc", ".docx"]:
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif file_path.suffix.lower() == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def index_standards_document(project_id: str, file_path: Path):
    """Index a standards document into a FAISS vector store."""
    logger.info(f"Indexing standards document for project: {project_id}, file: {file_path}")
    
    if not test_embedding_service():
        logger.error("Embedding service is not available")
        raise ValueError("Embedding service is not available")
    
    output_dir = STANDARDS_RAG_DIR / project_id / "faiss_index"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    text_content = extract_text_from_file(file_path)
    if not text_content.strip():
        logger.warning(f"No content extracted from {file_path}")
        return
    
    document = Document(
        page_content=text_content,
        metadata={
            "source": file_path.name,
            "type": f"standards_{file_path.suffix.lstrip('.')}",
            "project_id": project_id
        }
    )
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents([document])
    logger.info(f"Split standards document into {len(chunks)} chunks")
    
    try:
        vector_store = FAISS.load_local(
            str(output_dir),
            embedding_wrapper,
            allow_dangerous_deserialization=True
        )
        vector_store.add_documents(chunks)
        logger.info(f"Updated existing standards vector store for project: {project_id}")
    except Exception:
        vector_store = FAISS.from_documents(chunks, embedding_wrapper)
        logger.info(f"Created new standards vector store for project: {project_id}")
    
    try:
        vector_store.save_local(str(output_dir))
        logger.info(f"Saved standards vector store to {output_dir}")
    except Exception as e:
        logger.error(f"Error saving standards vector store: {str(e)}")
        raise
    
    metadata_path = output_dir.parent / "metadata.json"
    metadata = {
        "project_id": project_id,
        "total_documents": 1,
        "total_chunks": len(chunks),
        "embedding_model": AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
        "created_at": datetime.now().isoformat(),
    }
    
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2)
    
    logger.info(f"Standards document indexing completed for project: {project_id}")

def index_files_for_rag(project_id: str, cobol_json: Dict[str, Any]):
    """Index COBOL analysis JSON for RAG."""
    logger.info(f"Indexing files for RAG: {project_id}")
    
    if not test_embedding_service():
        logger.error("Embedding service is not available")
        raise ValueError("Embedding service is not available")
    
    output_dir = RAG_DIR / project_id / "faiss_index"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    documents = []
    files_data = cobol_json.get("files", [])
    programs = cobol_json.get("programs", {})
    
    if isinstance(files_data, list):
        logger.info(f"Found {len(files_data)} files to process (list format)")
        
        for file_data in files_data:
            if not isinstance(file_data, dict):
                continue
                
            file_name = file_data.get("file_name", "Unknown")
            file_type = file_data.get("file_type", "")
            
            logger.info(f"Processing file: {file_name} ({file_type})")
            
            content_parts = []
            content_parts.append(f"File: {file_name}")
            content_parts.append(f"Type: {file_type}")
            
            # Process divisions
            divisions = file_data.get("divisions", {})
            if divisions:
                # Identification division
                identification = divisions.get("identification", {})
                if identification:
                    program_id = identification.get("program_id")
                    if program_id:
                        content_parts.append(f"Program ID: {program_id}")
                
                # Data division
                data_div = divisions.get("data", {})
                if data_div:
                    working_storage = data_div.get("working_storage", [])
                    if working_storage:
                        content_parts.append("Working Storage:")
                        for var in working_storage:
                            if isinstance(var, dict):
                                var_details = f"  {var.get('name', '')}: level={var.get('level', '')}, type={var.get('type', '')}, picture={var.get('picture', '')}"
                                content_parts.append(var_details)
                    
                    linkage_section = data_div.get("linkage_section", [])
                    if linkage_section:
                        content_parts.append("Linkage Section:")
                        for var in linkage_section:
                            if isinstance(var, dict):
                                var_details = f"  {var.get('name', '')}: level={var.get('level', '')}, type={var.get('type', '')}, picture={var.get('picture', '')}"
                                content_parts.append(var_details)
                    
                    file_section = data_div.get("file_section", [])
                    if file_section:
                        content_parts.append("File Section:")
                        for var in file_section:
                            if isinstance(var, dict):
                                var_details = f"  {var.get('name', '')}: level={var.get('level', '')}, type={var.get('type', '')}, picture={var.get('picture', '')}"
                                content_parts.append(var_details)
                
                # Procedure division
                procedure_div = divisions.get("procedure", [])
                if procedure_div:
                    content_parts.append("Procedure Division:")
                    for proc in procedure_div:
                        if isinstance(proc, dict):
                            proc_details = f"  {proc.get('paragraph', '')}: {', '.join(proc.get('code', []))}"
                            content_parts.append(proc_details)
                        else:
                            content_parts.append(f"  {proc}")
            
            # Process CICS commands
            cics_commands = file_data.get("cics_commands", [])
            if cics_commands:
                content_parts.append("CICS Commands:")
                for cmd in cics_commands:
                    if isinstance(cmd, dict):
                        content_parts.append(f"  {cmd.get('command', '')}")
                    else:
                        content_parts.append(f"  {cmd}")
            
            # Process variables
            variables = file_data.get("variables", [])
            if variables:
                content_parts.append("Variables:")
                for var in variables:
                    if isinstance(var, dict):
                        var_details = f"  {var.get('name', '')}: level={var.get('level', '')}, type={var.get('type', '')}, picture={var.get('picture', '')}"
                        content_parts.append(var_details)
                    else:
                        content_parts.append(f"  {var}")
            
            # Process paragraphs
            paragraphs = file_data.get("paragraphs", [])
            if paragraphs:
                content_parts.append("Paragraphs:")
                for para in paragraphs:
                    if isinstance(para, dict):
                        para_details = f"  {para.get('paragraph', '')}: {', '.join(para.get('code', []))}"
                        content_parts.append(para_details)
                    else:
                        content_parts.append(f"  {para}")
            
            # Process copybooks
            copybooks = file_data.get("copybooks", [])
            if copybooks:
                content_parts.append("Copybooks:")
                for copybook in copybooks:
                    if isinstance(copybook, dict):
                        content_parts.append(f"  {copybook.get('name', '')}")
                    else:
                        content_parts.append(f"  {copybook}")
            
            # Process JCL definitions
            jcl_definitions = file_data.get("jcl_definitions")
            if jcl_definitions:
                content_parts.append("JCL Definitions:")
                for jcl in jcl_definitions:
                    if isinstance(jcl, dict):
                        content_parts.append(f"  {jcl.get('type', '')}: {jcl.get('details', '')}")
                    else:
                        content_parts.append(f"  {jcl}")
            
            if len(content_parts) > 2:  # More than just file name and type
                content = "\n".join(str(part) for part in content_parts if part)
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source": file_name,
                        "type": f"cobol_{file_type.lstrip('.')}" if file_type else "cobol_file",
                        "project_id": project_id
                    }
                ))
                logger.info(f"Added document for file: {file_name}")
    
    elif isinstance(files_data, dict):
        logger.info(f"Found {len(files_data)} files to process (dict format)")
        for program_name, program_data in files_data.items():
            content_parts = []
            
            logger.info(f"Processing program: {program_name}")
            logger.info(f"Program data keys: {list(program_data.keys()) if isinstance(program_data, dict) else 'Not a dict'}")
            
            content_parts.append(f"Program: {program_name}")
            
            if isinstance(program_data, dict):
                # Add description
                if program_data.get("description"):
                    content_parts.append(f"Description: {program_data['description']}")
                
                # Add working storage with details
                working_storage = program_data.get("working_storage", {})
                if not working_storage:
                    working_storage = program_data.get("variables", {})
                
                if working_storage:
                    content_parts.append("Working Storage/Variables:")
                    for var_name, var_info in working_storage.items():
                        if isinstance(var_info, dict):
                            var_details = f"  {var_name}: type={var_info.get('type', '')}, "
                            var_details += f"size={var_info.get('size', '')}, "
                            var_details += f"usage={var_info.get('usage', '')}"
                            content_parts.append(var_details.strip(", "))
                        else:
                            content_parts.append(f"  {var_name}: {var_info}")
                
                # Add procedures with details
                procedures = program_data.get("procedures", {})
                if not procedures:
                    procedures = program_data.get("paragraphs", {})
                    
                if procedures:
                    content_parts.append("Procedures/Paragraphs:")
                    for proc_name, proc_info in procedures.items():
                        proc_details = f"  {proc_name}: "
                        if isinstance(proc_info, dict):
                            proc_details += f"{proc_info.get('description', '')}"
                            if proc_info.get('code'):
                                proc_details += f"\n    Code: {proc_info['code']}"
                            if proc_info.get('content'):
                                proc_details += f"\n    Content: {proc_info['content']}"
                        else:
                            proc_details += str(proc_info)
                        content_parts.append(proc_details)
                
                # Add file operations
                file_operations = program_data.get("file_operations", [])
                if file_operations:
                    content_parts.append("File Operations:")
                    for op in file_operations:
                        content_parts.append(f"  {op}")
                
                # Add database operations
                db_operations = program_data.get("database_operations", [])
                if db_operations:
                    content_parts.append("Database Operations:")
                    for op in db_operations:
                        content_parts.append(f"  {op}")
                
                # Add CICS commands
                cics_commands = program_data.get("cics_commands", [])
                if cics_commands:
                    content_parts.append("CICS Commands:")
                    for cmd in cics_commands:
                        content_parts.append(f"  {cmd}")
                
                # Add any other relevant data
                for key, value in program_data.items():
                    if key not in ["description", "working_storage", "variables", "procedures", "paragraphs", 
                                 "file_operations", "database_operations", "cics_commands"]:
                        if value and str(value).strip():
                            content_parts.append(f"{key.replace('_', ' ').title()}: {value}")
            
            else:
                content_parts.append(f"Content: {program_data}")
            
            if len(content_parts) > 1:  # More than just the program name
                content = "\n".join(str(part) for part in content_parts if part)
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source": program_name,
                        "type": "cobol_program",
                        "project_id": project_id
                    }
                ))
                logger.info(f"Added document for program: {program_name}")
    
    # Process dependencies
    dependencies = cobol_json.get("dependencies", {})
    if dependencies:
        logger.info(f"Found dependencies to process")
        content_parts = ["Dependencies:"]
        
        if isinstance(dependencies, dict):
            for dep_name, dep_info in dependencies.items():
                content_parts.append(f"  {dep_name}: {dep_info}")
        else:
            content_parts.append(f"  {dependencies}")
        
        if len(content_parts) > 1:
            content = "\n".join(str(part) for part in content_parts if part)
            documents.append(Document(
                page_content=content,
                metadata={
                    "source": "project_dependencies",
                    "type": "dependencies",
                    "project_id": project_id
                }
            ))
            logger.info(f"Added document for dependencies")
    
    # Fallback extraction if no documents found
    if not documents:
        logger.warning("No documents found with standard keys. Attempting to extract any available content...")
        
        all_content = []
        
        def extract_text_content(obj, prefix=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if isinstance(value, (str, int, float)) and str(value).strip():
                        all_content.append(f"{prefix}{key}: {value}")
                    elif isinstance(value, (dict, list)):
                        extract_text_content(value, f"{prefix}{key}.")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    extract_text_content(item, f"{prefix}[{i}].")
            elif isinstance(obj, (str, int, float)) and str(obj).strip():
                all_content.append(f"{prefix}: {obj}")
        
        extract_text_content(cobol_json)
        
        if all_content:
            content = "\n".join(all_content)
            documents.append(Document(
                page_content=content,
                metadata={
                    "source": f"project_{project_id}",
                    "type": "cobol_analysis",
                    "project_id": project_id
                }
            ))
            logger.info(f"Created fallback document with {len(all_content)} content items")
    
    if not documents:
        logger.error("No documents found to index even after fallback extraction")
        raise ValueError("No content found in COBOL analysis to index")
    
    logger.info(f"Prepared {len(documents)} documents for indexing")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")
    
    try:
        vector_store = FAISS.load_local(
            str(output_dir),
            embedding_wrapper,
            allow_dangerous_deserialization=True
        )
        vector_store.add_documents(chunks)
        logger.info(f"Updated existing vector store for project: {project_id}")
    except Exception:
        vector_store = FAISS.from_documents(chunks, embedding_wrapper)
        logger.info(f"Created new vector store for project: {project_id}")
    
    try:
        vector_store.save_local(str(output_dir))
        logger.info(f"Saved vector store to {output_dir}")
    except Exception as e:
        logger.error(f"Error saving vector store: {str(e)}")
        raise
    
    metadata_path = output_dir.parent / "metadata.json"
    metadata = {
        "project_id": project_id,
        "total_documents": len(documents),
        "total_chunks": len(chunks),
        "embedding_model": AZURE_CONFIG["AZURE_OPENAI_EMBED_MODEL"],
        "created_at": datetime.now().isoformat(),
    }
    
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2)
    
    logger.info(f"RAG indexing completed successfully for project: {project_id}")

def load_vector_store(project_id: str):
    """Load and combine vector stores from both rag and standards-rag directories."""
    try:
        cobol_faiss_path = RAG_DIR / project_id / "faiss_index"
        standards_faiss_path = STANDARDS_RAG_DIR / project_id / "faiss_index"
        
        vector_stores = []
        
        if cobol_faiss_path.exists():
            cobol_vector_store = FAISS.load_local(
                str(cobol_faiss_path), 
                embedding_wrapper, 
                allow_dangerous_deserialization=True
            )
            vector_stores.append(cobol_vector_store)
            logger.info(f"COBOL vector store loaded successfully for project: {project_id}")
        
        if standards_faiss_path.exists():
            standards_vector_store = FAISS.load_local(
                str(standards_faiss_path),
                embedding_wrapper,
                allow_dangerous_deserialization=True
            )
            vector_stores.append(standards_vector_store)
            logger.info(f"Standards vector store loaded successfully for project: {project_id}")
        
        if not vector_stores:
            logger.warning(f"No vector stores found for project: {project_id}")
            return None
        
        if len(vector_stores) > 1:
            combined_vector_store = vector_stores[0]
            for vs in vector_stores[1:]:
                combined_vector_store.merge_from(vs)
            logger.info(f"Combined {len(vector_stores)} vector stores for project: {project_id}")
            return combined_vector_store
        else:
            return vector_stores[0]
        
    except Exception as e:
        logger.error(f"Error loading vector store for project {project_id}: {str(e)}")
        return None

def query_vector_store(vector_store, query: str, k: int = 3):
    """Query the combined vector store with similarity search."""
    try:
        if not vector_store:
            logger.warning("Vector store is None")
            return []
        
        logger.info(f"Performing similarity search with query: '{query}' and k={k}")
        results = vector_store.similarity_search(query, k=k)
        logger.info(f"Found {len(results)} results")
        
        for i, result in enumerate(results):
            logger.info(f"Result {i+1}: source={result.metadata.get('source', 'unknown')}, type={result.metadata.get('type', 'unknown')}")
            logger.info(f"Result {i+1} content preview: {result.page_content[:100]}...")
        
        return results
        
    except Exception as e:
        logger.error(f"Error querying vector store: {str(e)}")
        return []